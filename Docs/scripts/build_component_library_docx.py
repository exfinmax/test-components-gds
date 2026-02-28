from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt


def set_default_font(document: Document, font_name: str = "Arial", size: int = 11) -> None:
    style = document.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def main() -> None:
    root = Path(__file__).resolve().parents[2]
    output = root / "Docs" / "component_library_catalog_zh.docx"

    doc = Document()
    set_default_font(doc)

    doc.add_heading("ComponentLibrary 组件目录与接入说明", level=0)
    doc.add_paragraph(f"生成日期：{date.today().isoformat()}")

    add_heading(doc, "1. 架构分层")
    doc.add_paragraph("Global Services：全局依赖能力，仅保留一套（EventBus/ObjectPool/TimeController/LocalTimeDomain）。")
    doc.add_paragraph("Packs：按游戏品类拆分的可复制组件包。")
    doc.add_paragraph("Demos：每个品类都有最小可运行场景，用于快速验收。")
    doc.add_paragraph("Addon(可选)：component_library_share 仅做编辑器 Custom Type 注册。")

    add_heading(doc, "2. 统一依赖")
    dep_table = doc.add_table(rows=1, cols=2)
    dep_table.style = "Table Grid"
    dep_table.rows[0].cells[0].text = "文件"
    dep_table.rows[0].cells[1].text = "用途"
    deps = [
        ("ComponentLibrary/Dependencies/event_bus.gd", "全局事件总线"),
        ("ComponentLibrary/Dependencies/object_pool.gd", "全局对象池"),
        ("ComponentLibrary/Dependencies/time_controller.gd", "全局时间控制与冻结帧"),
        ("ComponentLibrary/Dependencies/local_time_domain.gd", "局部时间域父节点模板"),
        ("ComponentLibrary/Dependencies/component_base.gd", "通用组件基类"),
        ("ComponentLibrary/Dependencies/character_component_base.gd", "角色组件基类"),
    ]
    for path, usage in deps:
        row = dep_table.add_row().cells
        row[0].text = path
        row[1].text = usage

    add_heading(doc, "3. 品类组件包与演示")
    pack_table = doc.add_table(rows=1, cols=4)
    pack_table.style = "Table Grid"
    headers = ["品类", "核心组件", "模板场景", "Demo 场景"]
    for i, text in enumerate(headers):
        pack_table.rows[0].cells[i].text = text

    pack_rows = [
        ("Foundation", "DataBlackboardComponent 等", "Packs/Foundation/Templates/foundation_template.tscn", "Demos/Foundation/foundation_demo.tscn"),
        ("Action", "TriggerRouterComponent 等", "Packs/Action/Templates/action_template.tscn", "Demos/Action/action_demo.tscn"),
        ("Time", "TimelineSwitchComponent 等", "Packs/Time/Templates/time_template.tscn", "Demos/Time/time_demo.tscn"),
        ("UI", "UIPageStateComponent", "Packs/UI/Templates/ui_template.tscn", "Demos/UI/ui_demo.tscn"),
        ("VFX", "ImpactVFXComponent", "Packs/VFX/Templates/vfx_template.tscn", "Demos/VFX/vfx_demo.tscn"),
        ("Shooter", "ProjectileEmitterComponent", "Packs/Shooter/Templates/projectile_emitter_template.tscn", "Demos/Shooter/shooter_demo.tscn"),
        ("RPG", "AttributeSetComponent", "Packs/RPG/Templates/attribute_set_template.tscn", "Demos/RPG/rpg_demo.tscn"),
        ("Strategy", "ProductionQueueComponent", "Packs/Strategy/Templates/production_queue_template.tscn", "Demos/Strategy/strategy_demo.tscn"),
        ("Survival", "StatusEffectComponent", "Packs/Survival/Templates/status_effect_template.tscn", "Demos/Survival/survival_demo.tscn"),
        ("Card", "DeckDrawComponent", "Packs/Card/Templates/deck_draw_template.tscn", "Demos/Card/card_demo.tscn"),
        ("Puzzle", "SequenceSwitchComponent", "Packs/Puzzle/Templates/sequence_switch_template.tscn", "Demos/Puzzle/puzzle_demo.tscn"),
        ("Roguelike", "WeightedSpawnTableComponent", "Packs/Roguelike/Templates/weighted_spawn_table_template.tscn", "Demos/Roguelike/roguelike_demo.tscn"),
        ("Platformer", "CoyoteJumpComponent", "Packs/Platformer/Templates/coyote_jump_template.tscn", "Demos/Platformer/platformer_demo.tscn"),
        ("Racing", "LapCheckpointComponent", "Packs/Racing/Templates/lap_checkpoint_template.tscn", "Demos/Racing/racing_demo.tscn"),
        ("Builder", "GridPlacementComponent", "Packs/Builder/Templates/grid_placement_template.tscn", "Demos/Builder/builder_demo.tscn"),
    ]
    for genre, component, template, demo in pack_rows:
        row = pack_table.add_row().cells
        row[0].text = genre
        row[1].text = component
        row[2].text = f"ComponentLibrary/{template}"
        row[3].text = f"ComponentLibrary/{demo}"

    add_heading(doc, "4. 插件分发（可选）")
    doc.add_paragraph("插件目录：addons/component_library_share")
    doc.add_paragraph("作用：在编辑器节点创建面板中注册常用组件类型，提升发现效率。")
    doc.add_paragraph("边界：插件不接管运行时逻辑，不影响复制即用路径。")

    add_heading(doc, "5. 验收基线")
    doc.add_paragraph("每个品类包必须包含：")
    doc.add_paragraph("- Components/*.gd")
    doc.add_paragraph("- Templates/*.tscn")
    doc.add_paragraph("- Demos/<Genre>/*_demo.tscn + *_demo.gd")
    doc.add_paragraph("- README.md")

    doc.save(str(output))
    print(output)


if __name__ == "__main__":
    main()
